import os

import qrcode
from docx import Document
from docx.shared import Inches
from PIL import Image


def add_qr_to_docx(docx_path, qr_data, output_path, filename):
    # Создаем QR-код
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )

    # Добавляем данные с указанием кодировки UTF-8
    qr.add_data(qr_data.encode('utf-8'))  # Кодируем данные в UTF-8
    qr.make(fit=True)

    # Генерируем изображение QR-кода
    qr_img = qr.make_image(fill_color="black", back_color="white")
    qr_img_path = "qr_code.png"
    qr_img.save(qr_img_path)

    # Загружаем документ
    doc = Document(docx_path)

    # Вставляем QR-код в начало документа (перед первым параграфом)
    first_paragraph = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()  # Если параграфов нет, создаем новый
    run = first_paragraph.insert_paragraph_before().add_run()
    run.add_picture(qr_img_path, width=Inches(2))  # Указываем ширину в дюймах

    # Создаём директорию, если она не существует
    os.makedirs(output_path, exist_ok=True)

    # Сохраняем новый документ
    output_file_path = os.path.join(output_path, filename)
    doc.save(output_file_path)

    return output_file_path


